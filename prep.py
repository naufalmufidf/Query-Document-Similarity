import numpy as np
import pandas as pd
import sys
import PyPDF2 as pypd
import docx
import re
import string
import collections
import nltk
from nltk.tokenize import word_tokenize
# nltk.download('punkt')
from nltk.corpus import stopwords
# nltk.download('stopwords')
np.set_printoptions(threshold=sys.maxsize)
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import operator


def pdfview(path):
  openTxt = open(path, 'rb')
  readPDFTxt = pypd.PdfReader(openTxt)
  pages = len(readPDFTxt.pages)
  output = []
  for i in range(pages):
    page = readPDFTxt.pages[i]
    output.append(page.extract_text())

  seperator = ','
  PDFtxt = seperator.join(output)

  return PDFtxt

def wordview(path):
  openTxt = open(path, 'rb')
  txt = docx.Document(openTxt)
  Wordtxt = []
  for i in txt.paragraphs:
      Wordtxt.append(i.text)
  # return '\n'.join(Wordtxt)
  return Wordtxt

def pdf(path):
  openTxt = open(path, 'rb')
  readPDFTxt = pypd.PdfReader(openTxt)
  pages = len(readPDFTxt.pages)
  output = []
  for i in range(pages):
    page = readPDFTxt.pages[i]
    output.append(page.extract_text())
  openTxt.close()
  
  return output

def word(path):
  openTxt = open(path, 'rb')
  txt = docx.Document(openTxt)
  Wordtxt = []
  for i in txt.paragraphs:
      Wordtxt.append(i.text)
  openTxt.close()

  return Wordtxt

def lowercase(dataset):
  return dataset.apply(lambda x: ''.join([w for w in x.lower()]))

def preprocess_remove(to_remove, data):
    r = re.findall(to_remove, data)
    for i in r:
        data = re.sub(i, '', data)
    return data

from Sastrawi.Stemmer.StemmerFactory import StemmerFactory
factory = StemmerFactory()
stemmer = factory.create_stemmer()

def stem_sentences(sentence):
    tokens = sentence.split()
    stemmed_tokens = [stemmer.stem(token) for token in tokens]
    return ' '.join(stemmed_tokens)

def sw(data):
  stop_words = stopwords.words('indonesian')

  new_text = []
  word_dic = []
  for text in data:
      new_text2 = []
      for word in text.split():
          if word.lower() not in stop_words:
              new_text2.append(word.lower())
              word_dic.append(word.lower())
      new_text2 = " ".join(new_text2)
      new_text.append(new_text2)
  
  return new_text

def itung(data):
  total = len(data)
  print("Jumlah kata= ", total)
  
  count = collections.Counter(data).most_common()
  for i in count:
    print("%s\t: %d"%(i[0],i[1]))


def retrieve(path):
  doc = path
  if doc.endswith('.pdf'):
    hasil = pdf(doc)
    df = pd.DataFrame()
    df['Original'] = hasil
    df['Lower cased'] = lowercase(df['Original'])
    df['Clean_text'] = df['Lower cased']
    df['Clean_text'] = df['Clean_text'].str.replace('[^a-zA-Z]', ' ')
    df['Clean_text'] = np.vectorize(preprocess_remove)('http[\w]*',df['Clean_text'])
    df['Clean_text'] = np.vectorize(preprocess_remove)('@[\w]*',df['Clean_text'])
    df['Clean_text'] = df['Clean_text'].apply(lambda x: ' '.join([w for w in x.split() if len(w)>3]))
    df['Stemmed'] = df['Clean_text'].apply(stem_sentences)
    df['Final'] = sw(df['Stemmed'])
    temp = df['Final'].values.tolist()
    temp = str(temp)
    temp = temp.split()
    hasil = itung(temp)

  elif doc.endswith('.docx'):
    hasil = word(doc)
    df = pd.DataFrame()
    df['Original'] = hasil
    df['Lower cased'] = lowercase(df['Original'])
    df['Clean_text'] = df['Lower cased']
    df['Clean_text'] = df['Clean_text'].str.replace('[^a-zA-Z]', ' ')
    df['Clean_text'] = np.vectorize(preprocess_remove)('http[\w]*',df['Clean_text'])
    df['Clean_text'] = np.vectorize(preprocess_remove)('@[\w]*',df['Clean_text'])
    df['Clean_text'] = df['Clean_text'].apply(lambda x: ' '.join([w for w in x.split() if len(w)>3]))
    df['Stemmed'] = df['Clean_text'].apply(stem_sentences)
    df['Final'] = sw(df['Stemmed'])
    temp = df['Final'].values.tolist()
    temp = str(temp)
    temp = temp.split()
    hasil = itung(temp)

  else: 
    hasil = 'Harap masukan dokumen .pdf atau .docx'
  
  return print(hasil)

def prep(path):
  doc = path
  if doc.endswith('.pdf'):
    hasil = pdf(doc)
    df = pd.DataFrame()
    df['Original'] = hasil
    df['Lower cased'] = lowercase(df['Original'])
    df['Clean_text'] = df['Lower cased']
    df['Clean_text'] = df['Clean_text'].str.replace('[^a-zA-Z]', ' ')
    df['Clean_text'] = np.vectorize(preprocess_remove)('http[\w]*',df['Clean_text'])
    df['Clean_text'] = np.vectorize(preprocess_remove)('@[\w]*',df['Clean_text'])
    df['Clean_text'] = df['Clean_text'].apply(lambda x: ' '.join([w for w in x.split() if len(w)>3]))
    df['Stemmed'] = df['Clean_text'].apply(stem_sentences)
    df['Final'] = sw(df['Stemmed'])
    temp = df['Final'].values.tolist()
    temp = str(temp)
    temp = temp.split()
    # temp = ','.join(str(temp))
    hasil = temp
  
  elif doc.endswith('.docx'):
    hasil = word(doc)
    df = pd.DataFrame()
    df['Original'] = hasil
    df['Lower cased'] = lowercase(df['Original'])
    df['Clean_text'] = df['Lower cased']
    df['Clean_text'] = df['Clean_text'].str.replace('[^a-zA-Z]', ' ')
    df['Clean_text'] = np.vectorize(preprocess_remove)('http[\w]*',df['Clean_text'])
    df['Clean_text'] = np.vectorize(preprocess_remove)('@[\w]*',df['Clean_text'])
    df['Clean_text'] = df['Clean_text'].apply(lambda x: ' '.join([w for w in x.split() if len(w)>3]))
    df['Stemmed'] = df['Clean_text'].apply(stem_sentences)
    df['Final'] = sw(df['Stemmed'])
    temp = df['Final'].values.tolist()
    temp = str(temp)
    temp = temp.split()
    # temp = ','.join(str(temp))
    hasil = temp
  
  else: 
    hasil = 'Harap masukan dokumen .pdf atau .docx'
  
  return hasil

def prep2(path):
  doc = path
  if doc.endswith('.pdf'):
    hasil = pdf(doc)
    df = pd.DataFrame()
    df['Original'] = hasil
    df['Lower cased'] = lowercase(df['Original'])
    df['Clean_text'] = df['Lower cased']
    df['Clean_text'] = df['Clean_text'].str.replace('[^a-zA-Z]', ' ')
    df['Clean_text'] = np.vectorize(preprocess_remove)('http[\w]*',df['Clean_text'])
    df['Clean_text'] = np.vectorize(preprocess_remove)('@[\w]*',df['Clean_text'])
    df['Clean_text'] = df['Clean_text'].apply(lambda x: ' '.join([w for w in x.split() if len(w)>3]))
    df['Stemmed'] = df['Clean_text'].apply(stem_sentences)
    df['Final'] = sw(df['Stemmed'])
    temp = df['Final'].values.tolist()
    temp = str(temp)
    # temp = ','.join(str(temp))
    hasil = temp
  
  elif doc.endswith('.docx'):
    hasil = word(doc)
    df = pd.DataFrame()
    df['Original'] = hasil
    df['Lower cased'] = lowercase(df['Original'])
    df['Clean_text'] = df['Lower cased']
    df['Clean_text'] = df['Clean_text'].str.replace('[^a-zA-Z]', ' ')
    df['Clean_text'] = np.vectorize(preprocess_remove)('http[\w]*',df['Clean_text'])
    df['Clean_text'] = np.vectorize(preprocess_remove)('@[\w]*',df['Clean_text'])
    df['Clean_text'] = df['Clean_text'].apply(lambda x: ' '.join([w for w in x.split() if len(w)>3]))
    df['Stemmed'] = df['Clean_text'].apply(stem_sentences)
    df['Final'] = sw(df['Stemmed'])
    temp = df['Final'].values.tolist()
    temp = str(temp)
    # temp = ','.join(str(temp))
    hasil = temp
  
  else: 
    hasil = 'Harap masukan dokumen .pdf atau .docx'
  
  return hasil

'''
dataset = {}
dataset["d1"] = xyz
dataset["q"] = "hitung iterasi"

# print(prep)

def index(corpus):
  tf = CountVectorizer()
  term_doc_matrix = tf.fit_transform(dataset.values())
  df = pd.DataFrame(term_doc_matrix.toarray(), index=dataset.keys(), columns=tf.get_feature_names_out())
  df = df.transpose()
  
  return df

def weight(corpus):
  tfidf = TfidfVectorizer()
  weighted = tfidf.fit_transform(dataset.values())
  df2 = pd.DataFrame(weighted.toarray(), index=dataset.keys(), columns=tfidf.get_feature_names_out())
  df2.transpose()

cs = cosine_similarity(weighted, weighted)
df_cs = pd.DataFrame(cs, index=dataset.keys(), columns=dataset.keys())
df_cs

rank_cs = {}
for i in dataset.keys():
  if i != "q":
    rank_cs[i] = df_cs.at[i, "q"]
top_rank_cs = dict(sorted(rank_cs.items(), key=operator.itemgetter(1), reverse=True))
hasil = pd.DataFrame(top_rank_cs.values(), index=top_rank_cs.keys(), columns=["Cosine similarity"])
'''
